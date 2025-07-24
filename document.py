from docx import Document
import io
import re

def create_docx_in_memory(summary_json, document_title):
    doc = Document()
    doc.add_heading(document_title, level=0)

    for section in summary_json.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        for line in section["content"].split("\n"):
            line = line.strip()
            if line.startswith("- "):
                line = line[2:].strip()
                para = doc.add_paragraph(style="List Bullet")
                parts = re.split(r"(\*\*.*?\*\*)", line)
                for part in parts:
                    run = para.add_run()
                    if part.startswith("**") and part.endswith("**"):
                        run.text = part[2:-2]
                        run.bold = True
                    else:
                        run.text = part
            else:
                doc.add_paragraph(line.strip())
    
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream
